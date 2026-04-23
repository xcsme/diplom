#!/usr/bin/env python3
"""Convert the VKR markdown documentation to a properly formatted DOCX file."""

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
import re

doc = Document()

# Page setup
for section in doc.sections:
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(3)
    section.right_margin = Cm(1.5)

# ── Styles ──
style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(14)
style.paragraph_format.line_spacing = 1.5
style.paragraph_format.first_line_indent = Cm(1.25)

# Heading styles
for level in range(1, 4):
    hs = doc.styles[f'Heading {level}']
    hs.font.name = 'Times New Roman'
    hs.font.color.rgb = RGBColor(0, 0, 0)
    hs.font.bold = True
    hs.paragraph_format.first_line_indent = Cm(0)
    hs.paragraph_format.space_before = Pt(12)
    hs.paragraph_format.space_after = Pt(6)
    hs.paragraph_format.line_spacing = 1.5

doc.styles['Heading 1'].font.size = Pt(16)
doc.styles['Heading 2'].font.size = Pt(15)
doc.styles['Heading 3'].font.size = Pt(14)

def add_paragraph(text, bold=False, italic=False, align=None, font_size=None, first_indent=True):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(font_size or 14)
    run.bold = bold
    run.italic = italic
    if align == 'center':
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif align == 'right':
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    if not first_indent:
        p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.line_spacing = 1.5
    return p

def add_list_item(text, level=0):
    p = doc.add_paragraph(style='List Bullet')
    p.clear()
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(14)
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.left_indent = Cm(1.25 + level * 0.5)
    return p

def add_numbered_item(text, num):
    p = doc.add_paragraph()
    run = p.add_run(f'{num}. {text}')
    run.font.name = 'Times New Roman'
    run.font.size = Pt(14)
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.left_indent = Cm(1.25)
    return p

def add_code_block(lines):
    for line in lines:
        p = doc.add_paragraph()
        run = p.add_run(line)
        run.font.name = 'Courier New'
        run.font.size = Pt(10)
        p.paragraph_format.first_line_indent = Cm(0)
        p.paragraph_format.line_spacing = 1.0
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)

def add_table(headers, rows):
    table = doc.add_table(rows=len(rows) + 1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'
    # Header
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.font.name = 'Times New Roman'
                run.font.size = Pt(11)
                run.bold = True
    # Data
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = str(val)
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(11)
    doc.add_paragraph()  # spacing

# ══════════════════════════════════════════════════════
#  TITLE PAGE
# ══════════════════════════════════════════════════════

add_paragraph('МИНИСТЕРСТВО ЦИФРОВОГО РАЗВИТИЯ И\nМАССОВЫХ КОММУНИКАЦИЙ РОССИЙСКОЙ ФЕДЕРАЦИИ', align='center', font_size=12, first_indent=False)
add_paragraph('КОЛЛЕДЖ ТЕЛЕКОММУНИКАЦИЙ', align='center', font_size=12, first_indent=False)
add_paragraph('', first_indent=False)
add_paragraph('', first_indent=False)
add_paragraph('ДИПЛОМНЫЙ ПРОЕКТ', bold=True, align='center', font_size=16, first_indent=False)
add_paragraph('', first_indent=False)
add_paragraph('НА ТЕМУ:', align='center', font_size=14, first_indent=False)
add_paragraph('«Разработка веб-приложения экономического симулятора\n"Кофейня: Мастерская вкуса"»', bold=True, align='center', font_size=14, first_indent=False)
add_paragraph('', first_indent=False)
add_paragraph('Специальность: 09.02.07 Информационные системы и программирование', align='center', font_size=12, first_indent=False)
add_paragraph('', first_indent=False)
add_paragraph('', first_indent=False)
add_paragraph('', first_indent=False)
add_paragraph('', first_indent=False)
add_paragraph('', first_indent=False)
add_paragraph('', first_indent=False)
add_paragraph('г. Москва, 2026 г.', align='center', font_size=14, first_indent=False)

doc.add_page_break()

# ══════════════════════════════════════════════════════
#  Now parse the markdown and build the document
# ══════════════════════════════════════════════════════

with open('/app/docs/VKR_Diploma_Documentation.md', 'r') as f:
    content = f.read()

# Split by major sections using # headers
lines = content.split('\n')

in_code = False
code_lines = []
in_table = False
table_headers = []
table_rows = []
skip_title = True  # skip the first title block until we hit Введение

i = 0
while i < len(lines):
    line = lines[i]

    # Skip the initial title/TOC until we hit "# Введение"
    if skip_title:
        if line.strip() == '# Введение':
            skip_title = False
        else:
            i += 1
            continue

    # Code block handling
    if line.strip().startswith('```'):
        if in_code:
            add_code_block(code_lines)
            code_lines = []
            in_code = False
        else:
            in_code = True
        i += 1
        continue

    if in_code:
        code_lines.append(line)
        i += 1
        continue

    # Table handling
    if '|' in line and line.strip().startswith('|'):
        cells = [c.strip() for c in line.strip().split('|')[1:-1]]
        if all(set(c) <= set('-: ') for c in cells):
            # separator row, skip
            i += 1
            continue
        if not in_table:
            in_table = True
            table_headers = cells
        else:
            table_rows.append(cells)
        i += 1
        continue
    elif in_table:
        add_table(table_headers, table_rows)
        table_headers = []
        table_rows = []
        in_table = False

    stripped = line.strip()

    # Skip empty lines and horizontal rules
    if stripped == '' or stripped == '---':
        i += 1
        continue

    # Headings
    if stripped.startswith('# ') and not stripped.startswith('## '):
        title = stripped[2:].strip()
        doc.add_page_break()
        doc.add_heading(title, level=1)
        i += 1
        continue

    if stripped.startswith('## '):
        title = stripped[3:].strip()
        doc.add_heading(title, level=2)
        i += 1
        continue

    if stripped.startswith('### '):
        title = stripped[4:].strip()
        doc.add_heading(title, level=3)
        i += 1
        continue

    # Numbered list
    m = re.match(r'^(\d+)\.\s+(.+)$', stripped)
    if m:
        add_numbered_item(m.group(2), int(m.group(1)))
        i += 1
        continue

    # Bullet list
    if stripped.startswith('- '):
        text = stripped[2:]
        # Remove markdown bold
        text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)
        add_list_item(text)
        i += 1
        continue

    # Regular paragraph
    text = stripped
    # Remove markdown bold/italic for display
    text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)
    text = re.sub(r'\*(.+?)\*', r'\1', text)
    if text:
        add_paragraph(text)

    i += 1

# Flush remaining table
if in_table:
    add_table(table_headers, table_rows)

# Save
output_path = '/app/docs/VKR_Diploma_Кофейня_Мастерская_вкуса.docx'
doc.save(output_path)
print(f'Document saved to {output_path}')
print(f'Total paragraphs: {len(doc.paragraphs)}')
